#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from PIL import Image
from progressbar import progressbar
import time

image_path = '<path_to_your_image_file>'
text_path = '<path_to_your_text_file>'
output_path = '<output_path.docx>'

img = Image.open(image_path)
text = open(text_path).read()

width, height = img.size
print('Width: {}, height: {}'.format(width, height))
print('Total operations: {}'.format(width * height))

document = Document()
for section in document.sections:
  section.page_width = Inches(16.54)
  section.page_height = Inches(11.69)
styles = document.styles['Normal']
styles.font.size = Pt(1.5)
styles.font.name = 'Arial Black'
paragraph_format = styles.paragraph_format
paragraph_format.line_spacing = Pt(1)
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)

for i in progressbar(range(height)):
  p = document.add_paragraph()
  for j in range(width):
    r, g, b = img.getpixel((j, i))
    run = p.add_run(
      str(text[(j+i*width) % len(text)])
    )
    run.font.color.rgb = RGBColor(r, g, b)
    
document.save(output_path)